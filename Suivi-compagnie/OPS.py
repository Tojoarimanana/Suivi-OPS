import streamlit as st
import pandas as pd
import geopandas as gpd
import folium
from streamlit_folium import st_folium
import os
import tempfile
import zipfile
import io
import re
from docx import Document

# ────────────────────────────────────────────────
#  Fonctions utilitaires
# ────────────────────────────────────────────────

def format_date_fr(date):
    """Formate une date au format français : 15 mars 2024"""
    if pd.isna(date) or date is None:
        return "N/A"
    if isinstance(date, str):
        try:
            date = pd.to_datetime(date)
        except:
            return str(date)
    months = {
        1: 'janvier', 2: 'février', 3: 'mars', 4: 'avril',
        5: 'mai', 6: 'juin', 7: 'juillet', 8: 'août',
        9: 'septembre', 10: 'octobre', 11: 'novembre', 12: 'décembre'
    }
    try:
        return f"{date.day} {months[date.month]} {date.year}"
    except:
        return "N/A"


def clean_sheet_name(name):
    """Nettoie le nom d'onglet pour Excel (max 31 caractères, caractères interdits remplacés)"""
    name = re.sub(r'[\\/*?:\[\]]', '_', str(name))
    return name[:31] if name else "Onglet"


def export_to_excel(df_dict):
    """Exporte plusieurs dataframes dans un fichier Excel en mémoire"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for sheet_name, df in df_dict.items():
            safe_name = clean_sheet_name(sheet_name)
            df.to_excel(writer, sheet_name=safe_name, index=False)
    output.seek(0)
    return output.getvalue()


def export_to_word(df_dict):
    """Crée un document Word avec un tableau par dataframe"""
    doc = Document()
    for titre, df in df_dict.items():
        doc.add_heading(titre, level=2)
        if df.empty:
            doc.add_paragraph("Aucune donnée disponible.")
            continue

        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val) if pd.notna(val) else ""

        doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


# ────────────────────────────────────────────────
#  Chargement et formatage des données
# ────────────────────────────────────────────────

@st.cache_data(show_spinner="Chargement du fichier Excel...")
def load_data(file):
    df = pd.read_excel(file)

    date_columns = [
        'Date_de_signature_de_contrats', 'Date_d_entrée_en_vigeur',
        'Date_de_debut_de_la_phase', 'Date_de_la_fin_de_la_phase',
        'Date_du_dernier_MCM', 'Dernier_Paiement_de_frais_de_Formation',
        'Dernier_Paiement_de_frais_d_Administration', 'Dernier_Dépôt',
        'Date_de_Signature'
    ]

    for col in date_columns:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], errors='coerce')

    return df


@st.cache_data(show_spinner="Traitement du shapefile...")
def load_shapefile(zip_file):
    with tempfile.TemporaryDirectory() as tmpdir:
        zip_path = os.path.join(tmpdir, "shapefile.zip")
        with open(zip_path, "wb") as f:
            f.write(zip_file.getvalue())

        with zipfile.ZipFile(zip_path, "r") as zip_ref:
            zip_ref.extractall(tmpdir)

        shp_files = []
        for root, _, files in os.walk(tmpdir):
            for file in files:
                if file.lower().endswith(".shp"):
                    shp_files.append(os.path.join(root, file))

        if not shp_files:
            raise ValueError("Aucun fichier .shp trouvé dans l'archive ZIP.")

        return gpd.read_file(shp_files[0])


def format_df_for_display(df):
    df = df.copy()
    date_cols = [
        'Date_de_signature_de_contrats', 'Date_d_entrée_en_vigeur',
        'Date_de_debut_de_la_phase', 'Date_de_la_fin_de_la_phase',
        'Date_du_dernier_MCM', 'Dernier_Paiement_de_frais_de_Formation',
        'Dernier_Paiement_de_frais_d_Administration', 'Dernier_Dépôt',
        'Date_de_Signature'
    ]
    for col in date_cols:
        if col in df.columns:
            df[col] = df[col].apply(format_date_fr)
    return df


# ────────────────────────────────────────────────
#  Affichage carte
# ────────────────────────────────────────────────

def afficher_carte(df, gdf):
    if 'Nom' not in df.columns or 'Nom' not in gdf.columns:
        st.error("Colonne 'Nom' manquante dans les données ou le shapefile.")
        return

    noms_visibles = df['Nom'].dropna().unique()
    gdf_filt = gdf[gdf['Nom'].isin(noms_visibles)].copy()

    if gdf_filt.empty:
        st.warning("Aucun polygone ne correspond aux blocs filtrés.")
        return

    # Centre de la carte
    centroid = gdf_filt.geometry.centroid
    center_lat = centroid.y.mean()
    center_lon = centroid.x.mean()

    m = folium.Map(location=[center_lat, center_lon], zoom_start=6,
                   tiles="CartoDB positron")

    for _, row in gdf_filt.iterrows():
        nom = row.get('Nom', 'Inconnu')
        info = df[df['Nom'] == nom]
        if info.empty:
            continue
        info = info.iloc[0]

        popup_html = f"""
        <b>Bloc :</b> {nom}<br>
        <b>Compagnie :</b> {info.get('Compagnie', 'N/A')}<br>
        <b>Phase actuelle :</b> {info.get('Phases_actuelle', 'N/A')}<br>
        <b>Signature contrat :</b> {format_date_fr(info.get('Date_de_signature_de_contrats'))}<br>
        <b>Entrée en vigueur :</b> {format_date_fr(info.get('Date_d_entrée_en_vigeur'))}<br>
        <b>Début phase :</b> {format_date_fr(info.get('Date_de_debut_de_la_phase'))}<br>
        <b>Fin phase :</b> {format_date_fr(info.get('Date_de_la_fin_de_la_phase'))}<br>
        <hr>
        <b>Commentaires :</b> {info.get('Commentaires1', '—')}
        """

        folium.GeoJson(
            row['geometry'],
            name=nom,
            tooltip=folium.Tooltip(nom),
            popup=folium.Popup(popup_html, max_width=380),
            style_function=lambda x: {'fillColor': '#3388ff', 'color': '#0000ff', 'weight': 2, 'fillOpacity': 0.15}
        ).add_to(m)

    st.subheader("Carte des blocs pétroliers")
    st_folium(m, width=900, height=650, returned_objects=[])


# ────────────────────────────────────────────────
#  Affichage tableau avec colonnes choisies
# ────────────────────────────────────────────────

def afficher_tableau(df, colonnes, titre):
    if not colonnes:
        st.info("Aucune colonne sélectionnée.")
        return

    colonnes_existantes = [c for c in colonnes if c in df.columns]
    if not colonnes_existantes:
        st.warning("Aucune des colonnes sélectionnées n'existe dans les données.")
        return

    st.subheader(titre)
    df_show = format_df_for_display(df[colonnes_existantes])
    st.dataframe(df_show, use_container_width=True)


# ────────────────────────────────────────────────
#  APPLICATION PRINCIPALE
# ────────────────────────────────────────────────

def main():
    st.set_page_config(page_title="OMNIS – Suivi des Compagnies Pétrolières", layout="wide")

    st.title("🛢️ OMNIS – Suivi des Compagnies Pétrolières")

    # ── Chargement initial des fichiers ───────────────
    if 'raw_df' not in st.session_state or 'gdf' not in st.session_state:
        st.info("Veuillez charger les deux fichiers nécessaires pour commencer.")

        col1, col2 = st.columns(2)
        with col1:
            excel_file = st.file_uploader("Fichier Excel des contrats", type=["xlsx"], key="excel")
        with col2:
            zip_shp = st.file_uploader("Shapefile zippé (.zip)", type=["zip"], key="shp")

        if excel_file is not None:
            try:
                st.session_state.raw_df = load_data(excel_file)
                st.success("Fichier Excel chargé.")
                st.rerun()
            except Exception as e:
                st.error(f"Erreur lecture Excel : {e}")

        if zip_shp is not None:
            try:
                st.session_state.gdf = load_shapefile(zip_shp)
                st.success("Shapefile chargé.")
                st.rerun()
            except Exception as e:
                st.error(f"Erreur shapefile : {e}")

    else:
        st.success("Données et carte géographique chargées avec succès ✓")

        df = st.session_state.raw_df.copy()
        gdf = st.session_state.gdf

        # ── Filtre principal ───────────────────────────────
        compagnies = sorted(df['Compagnie'].dropna().unique())
        selected_comp = st.sidebar.selectbox(
            "Filtrer par compagnie",
            options=["Toutes"] + list(compagnies),
            index=0
        )

        if selected_comp != "Toutes":
            df = df[df['Compagnie'] == selected_comp]

        # ── Groupes de données / onglets ──────────────────
        groupes = {
            "Informations Compagnie / Bloc": [
                'Compagnie', 'Nom', 'Bloc', 'Coordonée_X', 'Coordonée_Y',
                'Date_de_signature_de_contrats', 'Date_d_entrée_en_vigeur'
            ],
            "Situation Actuelle": [
                'Phases_actuelle', 'Date_de_debut_de_la_phase',
                'Date_de_la_fin_de_la_phase', 'Situation_et_Activités_en_cours',
                'Travaux_déjà_réalisés', 'Commentaires1'
            ],
            "Termes Commerciaux": [
                'Cost_Recovery_Limit_(%)', 'Overhead_(%)',
                'Frais_d_Administration_(M_$)', 'Frais_de_Formation_(M_$)',
                'Bonus_de_Production_(M_$)',
                'Partage_de_Production_Pétrole_(Part_du_Gouvernement)',
                'Partage_de_Production_Gaz_(Part_du_Gouvernement)'
            ],
            "Obligations Contractuelles": [
                'Obligation_de_Travaux', 'Obligation_de_Rendu_(%)',
                'Obligation_de_Banque_Garantie_(M_$)', 'Travaux_réalisées',
                'Rendu_réalisé_(%)', 'Banque_Garantie_déposées_(M_$)',
                'Commentaires2'
            ],
            "MCM / TCM": [
                'Date_du_dernier_MCM', 'Lieu', 'Motifs', 'Résolution',
                'PTA_&_Budget', 'Réalisation_budgetaire', 'Commentaires3'
            ],
            "Obligations Financières": [
                'Frais_de_Formation', 'Dernier_Paiement_de_frais_de_Formation',
                'Frais_d_Administration', 'Dernier_Paiement_de_frais_d_Administration',
                'Garantie_Bancaire', 'Dernier_Dépôt', 'Observations'
            ],
            "Avenants": [
                'Dernier_Avenant', 'Date_de_Signature', 'Motifs_Avenant', 'Statut'
            ]
        }

        onglets = st.tabs(["Carte"] + list(groupes.keys()) + ["Rapport global"])

        # ── Onglet Carte ──────────────────────────────────
        with onglets[0]:
            afficher_carte(df, gdf)

        # ── Onglets par groupe ────────────────────────────
        colonnes_selectionnees = {}

        for idx, (titre, cols_defaut) in enumerate(groupes.items(), start=1):
            with onglets[idx]:
                st.subheader(f"→ {titre}")

                cols_dispo = [c for c in cols_defaut if c in df.columns]
                if not cols_dispo:
                    st.info("Aucune colonne de ce groupe présente dans le fichier.")
                    continue

                selection = st.multiselect(
                    "Colonnes à afficher",
                    options=cols_dispo,
                    default=cols_dispo,
                    key=f"sel_{titre}"
                )

                colonnes_selectionnees[titre] = selection

                if selection:
                    afficher_tableau(df, selection, titre)

        # ── Onglet Rapport global + export ────────────────
        with onglets[-1]:
            st.subheader("📄 Rapport récapitulatif")

            has_selection = any(cols for cols in colonnes_selectionnees.values())

            if not has_selection:
                st.info("Sélectionnez des colonnes dans les onglets précédents pour générer un rapport.")
            else:
                for titre, cols in colonnes_selectionnees.items():
                    if cols:
                        st.markdown(f"#### {titre}")
                        st.dataframe(format_df_for_display(df[cols]), use_container_width=True)

                format_export = st.radio("Format du rapport :", ["Excel (.xlsx)", "Word (.docx)"])

                if st.button("📥 Générer et télécharger le rapport"):
                    dfs_export = {}
                    for titre, cols in colonnes_selectionnees.items():
                        if cols:
                            cols_ok = [c for c in cols if c in df.columns]
                            if cols_ok:
                                dfs_export[titre] = format_df_for_display(df[cols_ok])

                    if not dfs_export:
                        st.warning("Aucune donnée valide à exporter.")
                    else:
                        if format_export == "Excel (.xlsx)":
                            data = export_to_excel(dfs_export)
                            st.download_button(
                                label="Télécharger rapport.xlsx",
                                data=data,
                                file_name="Suivi_Compagnies_Petroliferes.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        else:
                            data = export_to_word(dfs_export)
                            st.download_button(
                                label="Télécharger rapport.docx",
                                data=data,
                                file_name="Suivi_Compagnies_Petroliferes.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

    # ── Footer ────────────────────────────────────────────
    st.markdown("---")
    st.markdown("""
    <style>
    .footer-fixed {
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background-color: #f8f9fa;
        color: #444;
        text-align: center;
        padding: 10px 0;
        font-size: 0.95rem;
        border-top: 1px solid #dee2e6;
        z-index: 999;
    }
    </style>
    <div class="footer-fixed">
        <strong>Conçu par RANAIVOSOA Tojoarimanana Hiratriniala</strong>  
          —  Tél : +261 33 51 880 19
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
